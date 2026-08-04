#!/usr/bin/env python3
"""
img_probe.py
DOCX 안의 그림이 뷰어마다 안 보이는 원인을 한 번에 가려내는 시험 파일을 만든다.

왜 필요한가
    서버에서는 한/글도 워드도 띄울 수 없다. LibreOffice 로 그려 보면 잘 나오는데
    한/글에서는 그림이 하나도 안 보이고 워드에서는 잘려 보인다는 보고가 있었다.
    추측으로 고치면 왕복만 늘어난다. 그래서 그림 넣는 방식을 여러 가지로 만들어
    한 파일에 나란히 담는다. 한 번 열어 보면 어느 방식이 살아남는지 바로 보인다.

무엇을 가르는가
    가~라  그림을 넣는 마크업·형식이 문제인가
    마~사  2단 조판과 '문단 붙임 표시(keepNext/keepLines)' 가 문제인가

사용법:
    python3 scripts/img_probe.py --out out/그림시험.docx
"""

import argparse
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

EMU_MM = 36000


def pick_crops(crop_dir, n):
    """시험에 쓸 그림 몇 장을 고른다. 너무 작은 것은 판단이 어려우니 거른다."""
    names = sorted(x for x in os.listdir(crop_dir) if x.lower().endswith(".png"))
    out = []
    for name in names:
        p = os.path.join(crop_dir, name)
        if os.path.getsize(p) > 40_000:
            out.append(p)
        if len(out) >= n:
            break
    return out


def to_jpeg(png_path, out_dir):
    """PNG 를 JPEG 로 바꾼다. 뷰어가 PNG 만 못 읽는 경우를 가르기 위해서다."""
    try:
        from PIL import Image
    except ImportError:
        return None
    os.makedirs(out_dir, exist_ok=True)
    dst = os.path.join(out_dir, os.path.splitext(os.path.basename(png_path))[0] + ".jpg")
    Image.open(png_path).convert("RGB").save(dst, "JPEG", quality=92)
    return dst


MARK = "워드형태"     # '다/사' 방식을 저장 뒤에 골라내기 위한 표식


def move_ns_to_root(path):
    """저장한 파일을 열어 네임스페이스 선언 위치를 옮긴다.

    python-docx 는 그림마다 <wp:inline> 에 네임스페이스를 붙여 넣는다.
    문법상 맞지만 그걸 제대로 못 읽는 뷰어가 있다.
    워드가 실제로 쓰는 형태는 맨 위에 한 번만 선언하는 쪽이다.

    표식이 달린 그림에서만 붙어 있던 선언을 떼어 낸다.
    나머지는 그대로 둬야 두 방식을 견줄 수 있다.
    """
    import re
    import shutil
    import zipfile

    zin = zipfile.ZipFile(path)
    items = [(i, zin.read(i.filename)) for i in zin.infolist()]
    zin.close()

    def fix(xml):
        xml = xml.replace(
            "<w:document ",
            f'<w:document xmlns:a="{A_NS}" xmlns:pic="{PIC_NS}" ', 1)

        def one(m):
            blk = m.group()
            if MARK not in blk:
                return blk
            for decl in (f' xmlns:a="{A_NS}"', f' xmlns:pic="{PIC_NS}"'):
                blk = blk.replace(decl, "")
            return blk

        return re.sub(r"<wp:inline\b.*?</wp:inline>", one, xml, flags=re.S)

    tmp = path + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for info, data in items:
            if info.filename == "word/document.xml":
                data = fix(data.decode("utf-8")).encode("utf-8")
            zout.writestr(info, data)
    shutil.move(tmp, path)


def add_picture_plain(doc, path, width_mm, keep=True):
    """가 방식 — 지금 make_docx.py 가 쓰는 그대로."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = keep
    p.paragraph_format.keep_together = keep
    p.add_run().add_picture(path, width=Mm(width_mm))
    return p


def add_picture_wordlike(doc, path, width_mm, keep=True):
    """다 방식 — 워드가 실제로 쓰는 형태에 맞춘 마크업.

    네임스페이스는 맨 위에서 받아 쓰고, 여백 속성과 설명(descr)을 채운다.
    """
    p = add_picture_plain(doc, path, width_mm, keep)
    inline = p._p.findall(".//" + qn("wp:inline"))
    if not inline:
        return p
    el = inline[0]
    for a in ("distT", "distB", "distL", "distR"):
        el.set(a, "0")
    # 저장한 뒤에 이 표식을 보고 네임스페이스 선언을 떼어 낸다
    docpr = el.find(qn("wp:docPr"))
    if docpr is not None:
        docpr.set("descr", MARK)
    return p


def add_picture_in_table(doc, path, width_mm):
    """라 방식 — 표 한 칸 안에 넣는다. 표 안이면 살아나는 뷰어가 있다."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].add_run().add_picture(path, width=Mm(width_mm))
    return t


def note(doc, text, size=10, bold=False, before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p


def two_columns(section, on=True):
    sectPr = section._sectPr
    found = sectPr.xpath("./w:cols")
    cols = found[0] if found else sectPr.makeelement(qn("w:cols"), {})
    if not found:
        sectPr.append(cols)
    cols.set(qn("w:num"), "2" if on else "1")
    cols.set(qn("w:space"), "397")
    cols.set(qn("w:equalWidth"), "1")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--crops", default="./out/_crops")
    ap.add_argument("--out", default="./out/그림시험.docx")
    args = ap.parse_args()

    if not os.path.isdir(args.crops):
        sys.exit(f"오려낸 그림 폴더가 없습니다: {args.crops}")
    crops = pick_crops(args.crops, 6)
    if not crops:
        sys.exit("쓸 만한 그림을 찾지 못했습니다.")

    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    for a in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, a, Mm(14))
    two_columns(sec, on=False)

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(10)

    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = head.add_run("그림이 보이는 방식 찾기")
    r.bold = True
    r.font.size = Pt(14)

    note(doc, "각 항목 아래에 그림이 하나씩 있어야 합니다. "
              "보이는 것과 안 보이는 것을 알려 주세요.", size=9)

    note(doc, "가. 지금 방식 (PNG)", bold=True, before=14)
    add_picture_plain(doc, crops[0], 60)

    jpg = to_jpeg(crops[1], os.path.join(args.crops, "_jpg"))
    note(doc, "나. 같은 방식인데 그림만 JPEG", bold=True, before=14)
    if jpg:
        add_picture_plain(doc, jpg, 60)
    else:
        note(doc, "(pillow 가 없어 만들지 못함)", size=9)

    note(doc, "다. 워드가 쓰는 형태로 고친 마크업 (PNG)", bold=True, before=14)
    add_picture_wordlike(doc, crops[2], 60)

    note(doc, "라. 표 한 칸 안에 넣기 (PNG)", bold=True, before=14)
    add_picture_in_table(doc, crops[3], 60)

    # ── 2단 구역 ────────────────────────────────────────────────────
    new_sec = doc.add_section()
    new_sec.page_width = Mm(210)
    new_sec.page_height = Mm(297)
    for a in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(new_sec, a, Mm(14))
    two_columns(new_sec, on=True)

    note(doc, "여기부터는 실제 문제집과 같은 좌우 2단입니다.",
         bold=True, before=0)

    note(doc, "마. 2단 + 지금 방식 + 붙임 표시 있음", bold=True, before=12)
    add_picture_plain(doc, crops[4], 60, keep=True)

    note(doc, "바. 2단 + 지금 방식 + 붙임 표시 없음", bold=True, before=12)
    add_picture_plain(doc, crops[5 % len(crops)], 60, keep=False)

    note(doc, "사. 2단 + 워드 형태 마크업", bold=True, before=12)
    add_picture_wordlike(doc, crops[0], 60, keep=False)

    os.makedirs(os.path.dirname(os.path.abspath(args.out)) or ".", exist_ok=True)
    doc.save(args.out)
    move_ns_to_root(args.out)
    print(f"완료: {args.out}  (그림 {len(crops)}장 사용)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
