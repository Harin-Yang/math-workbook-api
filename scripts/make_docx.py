#!/usr/bin/env python3
"""
make_docx.py
추출한 문제를 좌우 2단 워드 문서로 조판한다. Mathpix 재호출 없음 = 비용 0원.

사용법:
    python3 scripts/make_docx.py --file 기하 --out out/기하.docx
    python3 scripts/make_docx.py --file 확률과통계 --out out/확통.docx --html out/확통.html

수식을 어떻게 넣는가
    Mathpix 가 준 LaTeX 를 워드 수식(OMML)으로 옮겨 넣는다. 한/글에서 편집된다.
    변환은 latex2omml.py 가 한다. 외부 패키지도 재호출도 없다.

    옮길 수 없는 문법을 만난 줄만 '원본에서 그 자리를 오려낸 그림' 으로 되돌린다.
    틀린 수식을 내보내느니 그림이 낫다.

줄을 어떻게 잇는가
    원본은 한 쪽을 두 단으로 짜 놓아 문장이 중간에서 끊긴다.
    그 끊김을 그대로 옮기면 우리 단 너비와 안 맞아 문장이 엉뚱한 데서 잘린다.
    Mathpix 가 줄마다 알려 주는 이어짐 표시로 문장을 다시 붙인다. (merge_lines)

크기를 어떻게 맞추는가
    오려낸 조각을 각자 원본 크기대로 넣으면 크기가 제각각이 된다.
    파일마다 '본문 한 줄이 가득 찼을 때의 폭' 을 배워 그걸 칼럼 폭에 맞춘다.
    원본에서의 상대 크기가 그대로 유지된다.

필요한 것
    pip install python-docx pymupdf
    원본 PDF 가 samples/테스트자료_스캔본/<run 폴더 이름>.pdf 에 있어야 한다.
"""

import argparse
import base64
import html
import os
import re
import sys
from collections import defaultdict

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import extract as EX  # noqa: E402
import latex2omml as L  # noqa: E402

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor
except ImportError:
    Document = None

# 수식이 들어 있는 줄인지
HAS_MATH = re.compile(r"\\\(|\\\[|\$|\\begin\{|\\frac|\\sqrt|\\overline|"
                      r"\\mathrm|\\vec|\\overrightarrow|\\angle|\\triangle")

# 이어지는 줄이 조사 한 글자로 시작하면 앞 낱말에 붙여야 한다.
# Mathpix 가 '변화' + '를 보이는' 을 띄어 이으라고 알려 주는 일이 있다.
# 조사 뒤에 공백이나 문장부호가 오는 경우만 잡는다. (한 글자 낱말 오인 방지)
JOSA_HEAD = re.compile(
    r"^(을|를|이|가|은|는|의|에|와|과|도|만|로|라|며|고|서|나|든)"
    r"(?=[\s,.)\]。，]|$)")

PAD = 6            # 오려낼 때 사방으로 남길 여백 (픽셀)
A4_W_MM = 210
MARGIN_MM = 14
GAP_MM = 7

# 문제와 문제 사이에 띄울 높이. 본문 10pt 기준 세 줄쯤 된다.
#
# 이 여백은 '앞 문제의 마지막 줄 뒤' 에 붙인다. 다음 문제 앞에 붙이지 않는다.
# 앞에 붙이면 문제가 다음 단으로 넘어갈 때 그 여백이 단 꼭대기에 그대로 남아
# 단마다 시작 높이가 달라진다. 뒤에 붙이면 단 끝에서 그냥 사라진다.
PROBLEM_GAP_PT = 36


def col_width_mm():
    return (A4_W_MM - MARGIN_MM * 2 - GAP_MM) / 2


def find_source_pdf(pdf_dirs, run_name):
    """run 폴더 이름과 같은 이름의 원본 PDF 를 찾는다."""
    for d in pdf_dirs:
        if not os.path.isdir(d):
            continue
        for ext in (".pdf", ".PDF"):
            p = os.path.join(d, run_name + ext)
            if os.path.exists(p):
                return p
        for n in os.listdir(d):
            if n.lower().endswith(".pdf") and os.path.splitext(n)[0] == run_name:
                return os.path.join(d, n)
    return None


def learn_body_width(rows):
    """파일마다 '본문 한 줄이 가득 찼을 때의 픽셀 폭' 을 배운다.

    오려낸 조각을 각자 원본 크기대로 넣으면 크기가 제각각이 된다.
    본문 폭을 칼럼 폭에 맞춰 두면 원본에서의 상대 크기가 그대로 유지된다.
    (원본에서 한 줄을 다 채우던 것은 칼럼도 다 채우고, 절반짜리는 절반)

    큰 그림이 폭을 부풀리지 않게 위쪽 10% 는 버리고 그 다음 값을 쓴다.
    """
    per = defaultdict(list)
    for ln in rows:
        if ln.get("type") in EX.FIGURE_TYPES:
            continue
        w = (ln.get("region") or {}).get("width")
        if isinstance(w, (int, float)) and w > 0:
            per[ln.get("_file")].append(w)
    out = {}
    for fname, ws in per.items():
        ws.sort()
        out[fname] = ws[int(len(ws) * 0.90)] if ws else None
    return out


def learn_page_widths(rows):
    """쪽마다 픽셀 폭을 정한다. 반환: {(파일, 쪽): 폭}

    Mathpix 가 page_width 를 주면 그걸 쓴다.
    없으면 그 쪽 줄들의 오른쪽 끝 최댓값으로 역산한다.
    """
    out, guess = {}, {}
    for ln in rows:
        key = (ln.get("_file"), ln.get("_page"))
        w = ln.get("_pw")
        if w:
            out[key] = w
            continue
        r = ln.get("region") or {}
        x, ww = r.get("top_left_x"), r.get("width")
        if isinstance(x, (int, float)) and isinstance(ww, (int, float)):
            guess[key] = max(guess.get(key, 0), x + ww)
    for key, v in guess.items():
        if key not in out:
            out[key] = v * 1.02
    return out


class Cropper:
    """원본 PDF 에서 필요한 사각형만 그려낸다."""

    def __init__(self, pdf_dirs, out_dir, body_widths=None):
        self.pdf_dirs = pdf_dirs
        self.out_dir = out_dir
        self.body_widths = body_widths or {}
        self.docs = {}       # run 이름 -> fitz.Document (없으면 None)
        self.pix = {}        # (run, page) -> (page, scale, 폭, 높이)
        self.n = 0
        self.fail = {}       # 실패 사유별 건수
        os.makedirs(out_dir, exist_ok=True)

    def note(self, why):
        self.fail[why] = self.fail.get(why, 0) + 1

    def doc_for(self, run_name):
        if run_name in self.docs:
            return self.docs[run_name]
        path = find_source_pdf(self.pdf_dirs, run_name)
        d = None
        if path and fitz:
            try:
                d = fitz.open(path)
            except Exception:
                d = None
        self.docs[run_name] = d
        return d

    def page_of(self, run_name, page_no, image_width):
        """(페이지, 배율, 픽셀폭, 픽셀높이). 못 얻으면 첫 값이 None.

        페이지 전체를 미리 그려두지 않는다. PyMuPDF 1.28 에서
        fitz.Pixmap(pixmap, IRect) 로 잘라내기가 깨져서,
        필요한 사각형만 clip 으로 바로 그리는 방식으로 바꿨다.
        """
        key = (run_name, page_no)
        if key in self.pix:
            return self.pix[key]
        d = self.doc_for(run_name)
        out = (None, 1.0, 0, 0)
        if d is not None and isinstance(page_no, int) \
                and 1 <= page_no <= d.page_count and image_width:
            pg = d[page_no - 1]
            scale = image_width / pg.rect.width
            out = (pg, scale,
                   int(round(pg.rect.width * scale)),
                   int(round(pg.rect.height * scale)))
        self.pix[key] = out
        return out

    def crop(self, ln, image_width, trim_left=0):
        """줄 하나를 오려 PNG 로 저장한다. 반환: (경로, 폭 mm) 또는 None

        trim_left: 왼쪽에서 이만큼(픽셀) 더 잘라낸다 — 문제 번호가 크롭에
        같이 찍히는 것을 막는 용도 (실물 피드백).
        """
        if fitz is None:
            self.note("pymupdf없음")
            return None
        r = ln.get("region") or {}
        x, y = r.get("top_left_x"), r.get("top_left_y")
        w, h = r.get("width"), r.get("height")
        if not all(isinstance(v, (int, float)) for v in (x, y, w, h)):
            self.note("좌표없음")
            return None
        if w <= 0 or h <= 0:
            self.note("크기0")
            return None
        if not image_width:
            self.note("쪽폭모름")
            return None

        pg, scale, pw_px, ph_px = self.page_of(ln.get("_file"),
                                               ln.get("_page"), image_width)
        if pg is None:
            self.note("원본쪽없음")
            return None

        x0 = max(0, int(x) - PAD + max(0, int(trim_left)))
        y0 = max(0, int(y) - PAD)
        x1 = min(pw_px, int(x + w) + PAD)
        y1 = min(ph_px, int(y + h) + PAD)
        if x1 <= x0 or y1 <= y0:
            self.note("범위밖")
            return None

        # 픽셀 사각형을 원본 점 단위로 되돌려 그 부분만 그린다
        clip = fitz.Rect(x0 / scale, y0 / scale, x1 / scale, y1 / scale)
        try:
            sub = pg.get_pixmap(matrix=fitz.Matrix(scale, scale), clip=clip)
        except Exception as e:
            self.note(f"자르기실패({type(e).__name__})")
            return None

        self.n += 1
        path = os.path.join(self.out_dir, f"crop_{self.n:05d}.png")
        try:
            sub.save(path)
        except Exception as e:
            self.note(f"저장실패({type(e).__name__})")
            return None

        # 크기 환산.
        # 본문 폭을 배웠으면 그걸 칼럼 폭에 맞춘다. 원본에서의 상대 크기가 유지된다.
        # 못 배웠으면 원본의 실제 물리 크기(점 -> mm)로 넣는다.
        body = self.body_widths.get(ln.get("_file"))
        if body:
            mm = (x1 - x0) / body * col_width_mm()
        else:
            mm = (x1 - x0) / scale * 25.4 / 72
        return path, mm


def set_two_columns(section, gap_mm=GAP_MM):
    """구역을 2단으로 만든다. python-docx 에 기능이 없어 XML 을 직접 건드린다."""
    sectPr = section._sectPr
    found = sectPr.xpath("./w:cols")
    if found:
        cols = found[0]
    else:
        cols = sectPr.makeelement(qn("w:cols"), {})
        sectPr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), str(int(gap_mm * 56.7)))   # mm -> twips
    cols.set(qn("w:equalWidth"), "1")


def add_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_text(doc, text, keep=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = keep
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_image(doc, path, width_mm, keep=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = keep
    run = p.add_run()
    w = min(width_mm, col_width_mm())
    try:
        run.add_picture(path, width=Mm(w))
    except Exception:
        run.add_text("[그림 넣기 실패]")
    return p


def add_rich(doc, parts, keep=True):
    """글자와 워드 수식이 섞인 한 줄. 수식은 한/글에서 편집된다."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = keep
    for part in parts:
        if part[0] == "t":
            run = p.add_run(part[1])
            run.font.size = Pt(10)
        else:
            p._p.append(parse_xml(part[1]))
    return p


def to_rich(text):
    """수식이 든 줄을 [글자 / 워드수식 / 브라우저수식] 조각으로 옮긴다.

    한 조각이라도 못 옮기면 None 을 준다. 부르는 쪽이 그림으로 되돌린다.
    틀린 수식을 내보내느니 그림이 낫다.
    """
    segs = L.segments(text)
    if not any(s[0] == "m" for s in segs):
        return None
    out = []
    for s in segs:
        if s[0] == "t":
            if s[1]:
                out.append(("t", s[1]))
            continue
        try:
            out.append(("m", L.to_omml(s[1]), L.to_mathml(s[1], s[2])))
        except (L.Unsupported, Exception):
            return None
    return out or None


def merge_lines(body):
    """원본 쪽에서 끊긴 줄을 문장 단위로 이어 붙인다.

    원본은 한 쪽에 두 단으로 짜여 있어 문장이 중간에서 끊긴다.
    그 끊김을 그대로 옮기면 우리 단 너비와 안 맞아 문장이 엉뚱한 데서 잘린다.

    Mathpix 가 줄마다 어떻게 끊겼는지 알려 준다.
        continues_line_space     앞줄에 이어짐 — 사이에 띄어쓰기
        continues_line_no_space  앞줄에 이어짐 — 사이를 띄우지 않는다
        continues_line_newline   진짜 줄바꿈 — 소문항 (1) (2) 같은 것
    앞의 둘만 이어 붙이고 나머지는 문단을 나눈다.

    반환: [{"lines": [줄...], "text": 이어붙인 글자} | {"fig": 줄}]
    """
    GLUE = {"continues_line_space": " ", "continues_line_no_space": ""}
    out = []
    follow_row = None      # 빈칸(□) 과 같은 시각적 줄 — 뒤따르는 글을 이어 붙인다
    for b in body:
        t = EX.txt(b)
        region = b.get("region") or {}
        if b.get("type") == "form_field":
            # 빈칸 상자 — text_display 가 '\begin{itemize}\item[□]' 같은 마크업으로
            # 와서 표 규칙(그림 되돌리기)을 오발시킨다 (실측: 동아 미적분 3쪽).
            # 글자 '□' 로 바꿔 앞 문장에 잇고, 같은 줄의 뒷글도 이어 붙인다.
            target = None
            for prev in reversed(out):
                if "text" in prev:
                    target = prev
                    break
                if "fig" not in prev:
                    break
            if target is not None:
                target["text"] = (target["text"] + " □").strip()
                target["lines"].append(b)
            else:
                out.append({"lines": [b], "text": "□"})
            y = region.get("top_left_y")
            h = region.get("height") or 0
            if isinstance(y, (int, float)):
                follow_row = (b.get("_page"), y - 10, y + h + 10)
            continue
        is_fig = b.get("type") in EX.FIGURE_TYPES or EX.is_image_md(t)
        glue = GLUE.get(b.get("subtype") or "")
        if follow_row is not None and not is_fig and glue is None and t:
            y = region.get("top_left_y")
            h = region.get("height") or 0
            if isinstance(y, (int, float)) and b.get("_page") == follow_row[0] \
                    and y < follow_row[2] and y + h > follow_row[1]:
                glue = " "          # '□ 안에 알맞은 것을' — 같은 줄이면 잇는다
        follow_row = None
        if not is_fig and glue is not None and out:
            # 잇는 대상은 마지막 '글' 단위 — 그림 줄이 문장 한가운데 끼어 있어도
            # 건너뛰고 잇는다. 그림에서 끊으면 '…도착했다고 / 한다' 처럼
            # 문단이 갈라진다 (실측: 벡터 18쪽 서술형 14).
            target = None
            for prev in reversed(out):
                if "text" in prev:
                    target = prev
                    break
                if "fig" not in prev:
                    break
            if target is not None:
                if glue and JOSA_HEAD.match(t):
                    glue = ""      # '변화' + '를 보이는' -> '변화를 보이는'
                target["text"] = (target["text"] + glue + t).strip()
                target["lines"].append(b)
                continue
        out.append({"fig": b} if is_fig else {"lines": [b], "text": t})
    return out


def iter_blocks(problems, page_widths, cropper):
    """문제들을 조판 조각으로 풀어 낸다.

    DOCX 와 HTML 미리보기가 이 하나를 같이 먹는다.
    두 벌로 쓰면 미리보기에서 통과한 조판이 DOCX 에서 달라진다.

    내보내는 조각:
        ("file",  파일명)
        ("label", 문제 머리)
        ("text",  글줄, 다음과 붙일지)
        ("image", PNG 경로, 폭mm, 다음과 붙일지)
        ("imgfail",)
        ("sep",)
    """
    cur_file = None
    for p in problems:
        if p["file"] != cur_file:
            cur_file = p["file"]
            yield ("file", str(cur_file))

        num = p["num"] if p["num"] is not None else ""
        yield ("label", f"{EX.display_name(p['name'])} {num}   (원본 {p['page']}쪽)")

        body = list(p["body"])
        for f in list(p["figs"]) + list(p["figs_guess"]):
            if f not in body:
                body.append(f)

        units = merge_lines(body)
        for k, u in enumerate(units):
            last = (k == len(units) - 1)

            # 진짜 그림은 오려낸다
            if "fig" in u:
                b = u["fig"]
                iw = page_widths.get((b.get("_file"), b.get("_page")))
                got = cropper.crop(b, iw)
                if got:
                    yield ("image", got[0], got[1], not last)
                else:
                    yield ("imgfail",)
                continue

            t = u["text"]
            if k == 0 and t:
                for _name, pat, _kind in EX.START_PATTERNS:
                    m = pat.match(t)
                    if m:
                        t = t[m.end():].strip()
                        break
                # '01. 원 …' 에서 번호만 떼면 '. 원 …' 이 남는다. 부스러기를 청소한다.
                t = t.lstrip(" .．)〕]")

            has_math = t and (HAS_MATH.search(t)
                              or any(EX.is_display_math(b)
                                     for b in u["lines"]))
            if has_math:
                # LaTeX 를 워드 수식으로 옮긴다.
                # 못 옮기면 그 줄만 예전처럼 그림으로 되돌린다.
                # 표(\begin{tabular})는 뼈대가 수식 구분자 밖에 있어 변환 검사를
                # 통과해 버린다 — 표가 보이면 무조건 그림으로 되돌린다.
                parts = None if ("\\begin{" in t and "\\begin{cases}" not in t) else to_rich(t)
                if parts:
                    yield ("rich", parts, not last)
                    continue
                for b in u["lines"]:
                    iw = page_widths.get((b.get("_file"), b.get("_page")))
                    got = cropper.crop(b, iw)
                    if got:
                        yield ("image", got[0], got[1], not last)
                    else:
                        yield ("imgfail",)
                        if EX.txt(b):
                            yield ("text", EX.txt(b), not last)
                continue

            if not t:
                continue
            yield ("text", t, not last)


def build(problems, page_widths, cropper, out_path, title):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Mm(MARGIN_MM)
    sec.right_margin = Mm(MARGIN_MM)
    sec.top_margin = Mm(14)
    sec.bottom_margin = Mm(14)
    set_two_columns(sec)

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(10)

    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = head.add_run(title)
    r.bold = True
    r.font.size = Pt(13)

    n_img, n_fail, n_math = 0, 0, 0

    # 문단을 만들기 전에 '다음에 무엇이 오는지' 를 먼저 본다.
    #   다음이 새 문제면  -> 여기서 여백을 주고 붙임을 없앱다
    #   아니면            -> 다음 줄과 붙여 둔다 (문제가 단에서 안 쪼개진다)
    events = [e for e in iter_blocks(problems, page_widths, cropper)]
    head.paragraph_format.space_after = Pt(8)

    for i, ev in enumerate(events):
        kind = ev[0]
        nxt = next((events[j][0] for j in range(i + 1, len(events))
                    if events[j][0] != "imgfail"), None)
        new_problem = nxt in ("label", "file")
        keep = not new_problem and nxt is not None

        if kind == "file":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            hr = p.add_run(ev[1])
            hr.bold = True
            hr.font.size = Pt(11)
        elif kind == "label":
            p = add_label(doc, ev[1])
        elif kind == "text":
            p = add_text(doc, ev[1], keep=keep)
        elif kind == "rich":
            p = add_rich(doc, ev[1], keep=keep)
            n_math += sum(1 for x in ev[1] if x[0] == "m")
        elif kind == "image":
            p = add_image(doc, ev[1], ev[2], keep=keep)
            n_img += 1
        else:                      # imgfail — 문단을 만들지 않는다
            n_fail += 1
            continue

        # 한 문단이 단 경계에서 반으로 쪼개지지 않게 한다
        p.paragraph_format.keep_together = True
        # 머리글은 늘 뒤 내용과 붙어 다닌다
        if kind in ("file", "label"):
            p.paragraph_format.keep_with_next = True
        if new_problem:
            p.paragraph_format.space_after = Pt(PROBLEM_GAP_PT)

    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".",
                exist_ok=True)
    doc.save(out_path)
    return n_img, n_fail, n_math


HTML_CSS = """
:root{
  --paper:#ffffff; --ink:#16181c; --ink-soft:#5a5f68;
  --rule:#d3cec6; --label:#2f5d8a;
  --desk:#ebe8e2; --chrome:#ffffff; --chrome-ink:#2a2d33;
  --chrome-rule:#dcd8d1;
}
@media (prefers-color-scheme: dark){
  :root{ --desk:#15171b; --chrome:#1c1f24; --chrome-ink:#dfe2e7;
         --chrome-rule:#2e333a; }
}
:root[data-theme="dark"]{ --desk:#15171b; --chrome:#1c1f24;
  --chrome-ink:#dfe2e7; --chrome-rule:#2e333a; }
:root[data-theme="light"]{ --desk:#ebe8e2; --chrome:#ffffff;
  --chrome-ink:#2a2d33; --chrome-rule:#dcd8d1; }

*{ box-sizing:border-box; }
body{
  margin:0; background:var(--desk);
  font-family:"Malgun Gothic","맑은 고딕","Apple SD Gothic Neo",
              "Noto Sans KR",sans-serif;
  color:var(--chrome-ink);
}
.bar{
  position:sticky; top:0; z-index:5;
  background:var(--chrome); border-bottom:1px solid var(--chrome-rule);
  padding:10px 18px; display:flex; gap:18px; align-items:baseline;
  flex-wrap:wrap; font-size:13px;
}
.bar b{ font-size:14px; }
.bar .n{ font-variant-numeric:tabular-nums; }
.bar .hint{ color:var(--ink-soft); margin-left:auto; }

/* 종이는 두 테마에서 모두 흰색이다. 인쇄물 미리보기라 일부러 그렇게 둔다. */
.sheet{
  width:210mm; margin:18px auto; padding:14mm;
  background:var(--paper); color:var(--ink);
  box-shadow:0 1px 3px rgba(0,0,0,.18);
}
.title{ text-align:center; font-weight:700; font-size:13pt;
        margin:0 0 10pt; text-wrap:balance; }
.cols{ column-count:2; column-gap:7mm;
       column-rule:1px solid var(--rule); }
.filehead{ font-weight:700; font-size:11pt; margin:0 0 4pt; }
/* 문제 사이 여백은 DOCX 와 같은 값(PROBLEM_GAP_PT)을 넣는다.
   앞 문제 뒤에 붙여야 단이 넘어갈 때 단 꼭대기에 여백이 남지 않는다. */
.q{ break-inside:avoid; page-break-inside:avoid; margin:0 0 __GAP__; }
.label{ font-weight:700; font-size:9pt; color:var(--label);
        margin:0 0 2pt; }
.line{ font-size:10pt; line-height:1.5; margin:0 0 2pt; }
.line img{ display:block; max-width:100%; margin:1pt 0; }
.miss{ font-size:9pt; color:#a8442a; }
/* 본문 속 분수도 교과서처럼 크게 편다. 기본값(compact)은 분수를 눌러 작게 그린다. */
math{ font-size:1.02em; math-style:normal; }
/* 선분 윗줄(overline) — 글꼴 기호는 안 늘어나 테두리 선으로 긋는다 */
.ovl{ border-top:.075em solid currentColor; padding-top:.1em; }
.unl{ border-bottom:.075em solid currentColor; padding-bottom:.1em; }
math[display="block"]{ display:block; margin:2pt 0; }

@media print{
  body{ background:#fff; }
  .bar{ display:none; }
  .sheet{ margin:0; box-shadow:none; width:auto; }
}
@media (max-width:230mm){
  .sheet{ width:auto; margin:12px; padding:8mm; }
}
"""


def build_html(problems, page_widths, cropper, out_path, title):
    """DOCX 와 같은 조각으로 브라우저 미리보기를 만든다.

    그림은 base64 로 파일 안에 박는다. 이 파일 하나만 열면 다 보인다.
    """
    out, n_img, n_fail, n_math = [], 0, 0, 0
    A = out.append

    A(f"<title>{html.escape(title)}</title>")
    A("<style>" + HTML_CSS.replace("__GAP__", f"{PROBLEM_GAP_PT}pt")
      + "</style>")
    A("<div class='bar'><b>" + html.escape(title) + "</b>")
    A(f"<span class='n'>문제 {len(problems)}개</span>")
    A("<span class='hint'>DOCX 와 같은 조판 코드로 그린 것이다. "
      "Ctrl+P 로 인쇄 모양까지 볼 수 있다.</span></div>")
    A("<div class='sheet'>")
    A(f"<div class='title'>{html.escape(title)}</div><div class='cols'>")

    open_q = False
    for ev in iter_blocks(problems, page_widths, cropper):
        kind = ev[0]
        if kind in ("file", "label") and open_q:
            A("</div>")
            open_q = False

        if kind == "file":
            A(f"<div class='filehead'>{html.escape(ev[1])}</div>")
        elif kind == "label":
            A("<div class='q'>")
            open_q = True
            A(f"<div class='label'>{html.escape(ev[1])}</div>")
        elif kind == "text":
            A(f"<div class='line'>{html.escape(ev[1])}</div>")
        elif kind == "rich":
            bits = []
            for part in ev[1]:
                if part[0] == "t":
                    bits.append(html.escape(part[1]))
                else:
                    bits.append(part[2])
                    n_math += 1
            A("<div class='line'>" + "".join(bits) + "</div>")
        elif kind == "image":
            try:
                with open(ev[1], "rb") as fh:
                    b64 = base64.b64encode(fh.read()).decode("ascii")
            except OSError:
                n_fail += 1
                continue
            w = min(ev[2], col_width_mm())
            A(f"<div class='line'><img src='data:image/png;base64,{b64}' "
              f"style='width:{w:.1f}mm' alt=''></div>")
            n_img += 1
        elif kind == "imgfail":
            n_fail += 1
            A("<div class='miss'>[그림 넣기 실패]</div>")

    if open_q:
        A("</div>")
    A("</div></div>")

    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".",
                exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(out))
    return n_img, n_fail, n_math


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--stage", default="./stage0_out")
    ap.add_argument("--file", action="append", default=[],
                    help="run 폴더 이름에 포함될 조각. 여러 번 줄 수 있다")
    ap.add_argument("--out", default="./out/문제집.docx")
    ap.add_argument("--title", default="추출 문제집")
    ap.add_argument("--pdfdir", action="append", default=[],
                    help="원본 PDF 폴더. 기본은 samples/테스트자료_스캔본 등")
    ap.add_argument("--cropdir", default="./out/_crops")
    ap.add_argument("--html", default=None,
                    help="같은 조판을 브라우저용 HTML 로도 저장한다")
    ap.add_argument("--html-only", action="store_true",
                    help="DOCX 는 만들지 않고 HTML 만")
    args = ap.parse_args()

    if args.html_only and not args.html:
        args.html = os.path.splitext(args.out)[0] + ".html"
    if Document is None and not args.html_only:
        sys.exit("python-docx 가 없습니다.  pip install python-docx pymupdf")
    if fitz is None:
        print("경고: pymupdf 가 없어 수식·그림을 넣을 수 없습니다.")

    runs = os.path.join(args.stage, "runs")
    if not os.path.isdir(runs):
        sys.exit(f"runs 폴더가 없습니다: {runs}")
    names = sorted(os.listdir(runs))
    if args.file:
        names = [n for n in names if any(f in n for f in args.file)]
    if not names:
        sys.exit("대상 run 폴더가 없습니다.")

    all_rows, pw = [], None
    for n in names:
        lj = os.path.join(runs, n, "result.lines.json")
        if not os.path.exists(lj):
            continue
        rows, w = EX.load(lj)
        for r in rows:
            r["_file"] = n
        all_rows.extend(rows)
        if w:
            pw = max(pw or 0, w)

    if not all_rows:
        sys.exit("lines.json 을 찾지 못했습니다.")

    page_widths = learn_page_widths(all_rows)
    problems, kept, dropped, live, st = EX.build(all_rows, pw)

    pdf_dirs = args.pdfdir or [
        "samples/테스트자료_스캔본",
        "samples/테스트자료_텍스트레이어",
        "s_scan", "s_text",
    ]
    body_widths = learn_body_width(all_rows)
    cropper = Cropper(pdf_dirs, args.cropdir, body_widths)

    if args.html_only:
        n_img, n_fail, n_math = 0, 0, 0
    else:
        n_img, n_fail, n_math = build(problems, page_widths, cropper,
                                      args.out, args.title)
        print(f"완료: {args.out}")

    if args.html:
        h = build_html(problems, page_widths, cropper, args.html,
                       args.title)
        print(f"완료: {args.html}  (브라우저로 열면 된다)")
        if args.html_only:
            n_img, n_fail, n_math = h

    print(f"문제 {len(problems)}개 / 편집 가능한 수식 {n_math}개 / "
          f"오려낸 그림 {n_img}장 / 실패 {n_fail}건")
    if cropper.fail:
        print("실패 사유:", ", ".join(f"{k} {v}" for k, v in
                                    sorted(cropper.fail.items())))
    missing = [k for k, v in cropper.docs.items() if v is None]
    if missing:
        print("원본 PDF 를 못 찾은 파일:")
        for m in missing:
            print(f"  {m}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
