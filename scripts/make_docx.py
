#!/usr/bin/env python3
"""
make_docx.py
추출한 문제를 좌우 2단 워드 문서로 조판한다. Mathpix 재호출 없음 = 비용 0원.

사용법:
    python3 scripts/make_docx.py --file 기하 --out out/기하.docx
    python3 scripts/make_docx.py --file 확률과통계 --out out/확통.docx --answer-lines 5

수식을 어떻게 넣는가
    지금은 '원본 페이지에서 그 자리를 오려낸 그림' 으로 넣는다.
    Mathpix 가 모든 줄의 좌표를 주므로 원본 PDF 를 같은 해상도로 그려서
    그 사각형만 잘라내면 된다. 화면에 보이던 그대로 나온다.

    대신 그 줄은 편집할 수 없다. 조판이 제대로 되는지 먼저 눈으로 확인하는 게
    목적이므로 이 방식으로 시작한다. 수식을 한/글에서 고칠 수 있어야 하면
    LaTeX -> OMML 변환기를 붙여야 하는데, 그건 조판이 통과된 뒤에 한다.

    수식이 없는 줄은 진짜 글자로 들어가므로 편집된다.

필요한 것
    pip install python-docx pymupdf
    원본 PDF 가 samples/테스트자료_스캔본/<run 폴더 이름>.pdf 에 있어야 한다.
    없으면 그 줄은 글자로 대체하고 넘어간다.
"""

import argparse
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import extract as EX  # noqa: E402

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor
except ImportError:
    Document = None

# 수식이 들어 있는 줄인지
HAS_MATH = re.compile(r"\\\(|\\\[|\$|\\begin\{|\\frac|\\sqrt|\\overline|"
                      r"\\mathrm|\\vec|\\overrightarrow|\\angle|\\triangle")

PAD = 6            # 오려낼 때 사방으로 남길 여백 (픽셀)
A4_W_MM = 210
MARGIN_MM = 14
GAP_MM = 7


def col_width_mm():
    return (A4_W_MM - MARGIN_MM * 2 - GAP_MM) / 2


def find_source_pdf(pdf_dirs, run_name):
    """run 폴더 이름과 같은 이름의 원본 PDF 를 찾는다."""
    for d in pdf_dirs:
        if not os.path.isdir(d):
            continue
        for ext in (".pdf", ".PDF"):
            p = os.path.join(d, run_name + ext)
            if os.path.exists(p):
                return p
        for n in os.listdir(d):
            if n.lower().endswith(".pdf") and os.path.splitext(n)[0] == run_name:
                return os.path.join(d, n)
    return None


def learn_page_widths(rows):
    """쪽마다 픽셀 폭을 정한다. 반환: {(파일, 쪽): 폭}

    Mathpix 가 page_width 를 주면 그걸 쓴다.
    없으면 그 쪽 줄들의 오른쪽 끝 최댓값으로 역산한다.
    머리말·꼬리말이 여백 가까이까지 묻어 있어 실제 폭에 거의 붙는다.
    """
    out, guess = {}, {}
    for ln in rows:
        key = (ln.get("_file"), ln.get("_page"))
        w = ln.get("_pw")
        if w:
            out[key] = w
            continue
        r = ln.get("region") or {}
        x, ww = r.get("top_left_x"), r.get("width")
        if isinstance(x, (int, float)) and isinstance(ww, (int, float)):
            guess[key] = max(guess.get(key, 0), x + ww)
    for key, v in guess.items():
        if key not in out:
            out[key] = v * 1.02      # 오른쪽 여백 몲을 조금 엹는다
    return out


class Cropper:
    """원본 PDF 페이지를 Mathpix 와 같은 해상도로 그려두고 잘라 쓴다."""

    def __init__(self, pdf_dirs, out_dir):
        self.pdf_dirs = pdf_dirs
        self.out_dir = out_dir
        self.docs = {}       # run 이름 -> fitz.Document (없으면 None)
        self.pix = {}        # (run, page) -> (pixmap, scale)
        self.n = 0
        self.fail = {}       # 실패 사유별 건수
        os.makedirs(out_dir, exist_ok=True)

    def note(self, why):
        self.fail[why] = self.fail.get(why, 0) + 1

    def doc_for(self, run_name):
        if run_name in self.docs:
            return self.docs[run_name]
        path = find_source_pdf(self.pdf_dirs, run_name)
        d = None
        if path and fitz:
            try:
                d = fitz.open(path)
            except Exception:
                d = None
        self.docs[run_name] = d
        return d

    def page_pixmap(self, run_name, page_no, image_width):
        key = (run_name, page_no)
        if key in self.pix:
            return self.pix[key]
        d = self.doc_for(run_name)
        out = (None, 1.0)
        if d is not None and isinstance(page_no, int) \
                and 1 <= page_no <= d.page_count and image_width:
            pg = d[page_no - 1]
            scale = image_width / pg.rect.width
            try:
                out = (pg.get_pixmap(matrix=fitz.Matrix(scale, scale)), scale)
            except Exception:
                out = (None, 1.0)
        self.pix[key] = out
        return out

    def crop(self, ln, image_width):
        """줄 하나를 오려 PNG 로 저장한다. 반환: (경로, 폭 mm) 또는 None"""
        if fitz is None:
            self.note("pymupdf없음")
            return None
        r = ln.get("region") or {}
        x, y = r.get("top_left_x"), r.get("top_left_y")
        w, h = r.get("width"), r.get("height")
        if not all(isinstance(v, (int, float)) for v in (x, y, w, h)):
            self.note("좌표없음")
            return None
        if w <= 0 or h <= 0:
            self.note("크기0")
            return None
        if not image_width:
            self.note("쪽폭모름")
            return None

        pm, scale = self.page_pixmap(ln.get("_file"), ln.get("_page"),
                                     image_width)
        if pm is None:
            self.note("원본쪽없음")
            return None

        x0 = max(0, int(x) - PAD)
        y0 = max(0, int(y) - PAD)
        x1 = min(pm.width, int(x + w) + PAD)
        y1 = min(pm.height, int(y + h) + PAD)
        if x1 <= x0 or y1 <= y0:
            self.note("범위밖")
            return None

        try:
            sub = fitz.Pixmap(pm, fitz.IRect(x0, y0, x1, y1))
        except Exception:
            self.note("자르기실패")
            return None

        self.n += 1
        path = os.path.join(self.out_dir, f"crop_{self.n:05d}.png")
        try:
            sub.save(path)
        except Exception:
            self.note("저장실패")
            return None

        # 픽셀 -> mm (원본 PDF 의 점 단위를 거쳐 환산)
        mm = (x1 - x0) / scale * 25.4 / 72
        return path, mm


def set_two_columns(section, gap_mm=GAP_MM):
    """구역을 2단으로 만든다. python-docx 에 기능이 없어 XML 을 직접 건드린다."""
    sectPr = section._sectPr
    found = sectPr.xpath("./w:cols")
    if found:
        cols = found[0]
    else:
        cols = sectPr.makeelement(qn("w:cols"), {})
        sectPr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), str(int(gap_mm * 56.7)))   # mm -> twips
    cols.set(qn("w:equalWidth"), "1")


def add_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_text(doc, text, keep=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = keep
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_image(doc, path, width_mm, keep=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = keep
    run = p.add_run()
    w = min(width_mm, col_width_mm())
    try:
        run.add_picture(path, width=Mm(w))
    except Exception:
        run.add_text("[그림 넣기 실패]")
    return p


def add_answer_space(doc, lines):
    for k in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = (k < lines - 1)
        p.add_run(" ").font.size = Pt(10)


def build(problems, page_widths, cropper, out_path, title, answer_lines):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Mm(MARGIN_MM)
    sec.right_margin = Mm(MARGIN_MM)
    sec.top_margin = Mm(14)
    sec.bottom_margin = Mm(14)
    set_two_columns(sec)

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(10)

    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = head.add_run(title)
    r.bold = True
    r.font.size = Pt(13)

    cur_file = None
    n_img, n_fail = 0, 0

    for p in problems:
        if p["file"] != cur_file:
            cur_file = p["file"]
            hp = doc.add_paragraph()
            hp.paragraph_format.space_before = Pt(12)
            hp.paragraph_format.keep_with_next = True
            hr = hp.add_run(str(cur_file))
            hr.bold = True
            hr.font.size = Pt(11)

        num = p["num"] if p["num"] is not None else ""
        add_label(doc, f"{p['name']} {num}   (원본 {p['page']}쪽)")

        body = list(p["body"])
        for f in list(p["figs"]) + list(p["figs_guess"]):
            if f not in body:
                body.append(f)

        for k, b in enumerate(body):
            last = (k == len(body) - 1)
            t = EX.txt(b)
            typ = b.get("type")
            iw = page_widths.get((b.get("_file"), b.get("_page")))

            if typ in EX.FIGURE_TYPES or EX.is_image_md(t) or \
                    (t and HAS_MATH.search(t)) or EX.is_display_math(b):
                got = cropper.crop(b, iw)
                if got:
                    add_image(doc, got[0], got[1], keep=not last)
                    n_img += 1
                else:
                    n_fail += 1
                    if t and not EX.is_image_md(t):
                        add_text(doc, t, keep=not last)
                continue

            if not t:
                continue
            if k == 0:
                for _name, pat, _kind in EX.START_PATTERNS:
                    m = pat.match(t)
                    if m:
                        t = t[m.end():].strip()
                        break
                if not t:
                    continue
            add_text(doc, t, keep=not last)

        add_answer_space(doc, answer_lines)

    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".",
                exist_ok=True)
    doc.save(out_path)
    return n_img, n_fail


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--stage", default="./stage0_out")
    ap.add_argument("--file", action="append", default=[],
                    help="run 폴더 이름에 포함될 조각. 여러 번 줄 수 있다")
    ap.add_argument("--out", default="./out/문제집.docx")
    ap.add_argument("--title", default="추출 문제집")
    ap.add_argument("--answer-lines", type=int, default=3,
                    help="문제마다 답 쓸 빈 줄 수")
    ap.add_argument("--pdfdir", action="append", default=[],
                    help="원본 PDF 폴더. 기본은 samples/테스트자료_스캔본 등")
    ap.add_argument("--cropdir", default="./out/_crops")
    args = ap.parse_args()

    if Document is None:
        sys.exit("python-docx 가 없습니다.  pip install python-docx pymupdf")
    if fitz is None:
        print("경고: pymupdf 가 없어 수식·그림을 넣을 수 없습니다.")
        print("      pip install pymupdf 후 다시 실행하세요.")

    runs = os.path.join(args.stage, "runs")
    if not os.path.isdir(runs):
        sys.exit(f"runs 폴더가 없습니다: {runs}")
    names = sorted(os.listdir(runs))
    if args.file:
        names = [n for n in names if any(f in n for f in args.file)]
    if not names:
        sys.exit("대상 run 폴더가 없습니다.")

    all_rows, pw = [], None
    for n in names:
        lj = os.path.join(runs, n, "result.lines.json")
        if not os.path.exists(lj):
            continue
        rows, w = EX.load(lj)
        for r in rows:
            r["_file"] = n
        all_rows.extend(rows)
        if w:
            pw = max(pw or 0, w)

    if not all_rows:
        sys.exit("lines.json 을 찾지 못했습니다.")

    page_widths = learn_page_widths(all_rows)
    problems, kept, dropped, live, st = EX.build(all_rows, pw)

    pdf_dirs = args.pdfdir or [
        "samples/테스트자료_스캔본",
        "samples/테스트자료_텍스트레이어",
        "s_scan", "s_text",
    ]
    cropper = Cropper(pdf_dirs, args.cropdir)

    n_img, n_fail = build(problems, page_widths, cropper, args.out,
                          args.title, args.answer_lines)

    print(f"완료: {args.out}")
    print(f"문제 {len(problems)}개 / 오려낸 그림 {n_img}장 / 실패 {n_fail}건")
    if cropper.fail:
        print("실패 사유:", ", ".join(f"{k} {v}" for k, v in
                                    sorted(cropper.fail.items())))
    missing = [k for k, v in cropper.docs.items() if v is None]
    if missing:
        print("원본 PDF 를 못 찾은 파일:")
        for m in missing:
            print(f"  {m}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
